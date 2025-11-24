"""File operation tools for reading code and writing documents"""

import os
import ast
from pathlib import Path
from typing import List, Dict, Optional, Set
import asyncio
import aiofiles

# Optional libraries for document conversion
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None

try:
    from markdown import markdown as markdown_to_html
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown_to_html = None

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    BeautifulSoup = None

try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError, Exception) as e:
    # WeasyPrint requires external libraries (GTK, Pango, etc.)
    # Gracefully degrade if not available
    WEASYPRINT_AVAILABLE = False
    HTML = None

# File size limits (bytes)
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB default
MAX_TOTAL_SIZE = 100 * 1024 * 1024  # 100 MB total


class FileTools:
    """Tools for file system operations"""
    
    @staticmethod
    def read_code_files(
        directory: str,
        extensions: Optional[List[str]] = None,
        max_file_size: int = MAX_FILE_SIZE,
        max_total_size: int = MAX_TOTAL_SIZE
    ) -> Dict[str, str]:
        """Read all code files from a directory with memory safety.
        
        Args:
            directory: Path to the directory
            extensions: List of file extensions to include (e.g., ['.py', '.java'])
                       If None, reads common code extensions
            max_file_size: Maximum size per file in bytes (default 10MB)
            max_total_size: Maximum total size in bytes (default 100MB)
        
        Returns:
            Dictionary mapping file paths to their contents
        """
        if extensions is None:
            extensions = ['.py', '.java', '.js', '.ts', '.cs', '.cpp', '.c', '.h', '.go', '.rs']
        
        code_files = {}
        directory_path = Path(directory)
        total_size = 0
        skipped_files = []
        
        if not directory_path.exists():
            return {"error": f"Directory {directory} does not exist"}
        
        for ext in extensions:
            for file_path in directory_path.rglob(f"*{ext}"):
                try:
                    # Check file size before reading
                    file_size = file_path.stat().st_size
                    
                    if file_size > max_file_size:
                        skipped_files.append(f"{file_path.name} (too large: {file_size / 1024 / 1024:.2f} MB)")
                        continue
                    
                    if total_size + file_size > max_total_size:
                        skipped_files.append(f"{file_path.name} (total size limit reached)")
                        continue
                    
                    with open(file_path, 'r', encoding='utf-8') as f:
                        relative_path = file_path.relative_to(directory_path)
                        code_files[str(relative_path)] = f.read()
                        total_size += file_size
                        
                except UnicodeDecodeError:
                    code_files[str(file_path)] = "[Binary file - skipped]"
                except Exception as e:
                    code_files[str(file_path)] = f"Error reading file: {str(e)}"
        
        # Add metadata about skipped files
        if skipped_files:
            code_files["_metadata_skipped"] = "\n".join([
                "Files skipped due to size limits:",
                *skipped_files,
                f"\nTotal size read: {total_size / 1024 / 1024:.2f} MB"
            ])
        
        return code_files
    
    @staticmethod
    def read_single_file(file_path: str) -> str:
        """
        Read a single file
        
        Args:
            file_path: Path to the file
        
        Returns:
            File contents as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            return f"Error reading file {file_path}: {str(e)}"
    
    @staticmethod
    async def write_document_async(content: str, output_path: str, format: str = "md") -> str:
        """Async write document to file using aiofiles.
        
        Args:
            content: Document content
            output_path: Output file path
            format: Output format (md, txt)
        
        Returns:
            Success message or error
        """
        try:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            async with aiofiles.open(output_file, 'w', encoding='utf-8') as f:
                await f.write(content)
            
            return f"Document successfully written to {output_path}"
        except Exception as e:
            return f"Error writing document: {str(e)}"
    
    @staticmethod
    def write_document(content: str, output_path: str, format: str = "md") -> str:
        """Write document to file (synchronous version).
        
        For async operations, use write_document_async instead.
        
        Args:
            content: Document content
            output_path: Output file path
            format: Output format (md, txt)
        
        Returns:
            Success message or error
        """
        try:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"Document successfully written to {output_path}"
        except Exception as e:
            return f"Error writing document: {str(e)}"
    
    @staticmethod
    def list_files(directory: str, pattern: str = "*") -> List[str]:
        """
        List files in a directory matching a pattern
        
        Args:
            directory: Directory path
            pattern: Glob pattern (e.g., "*.py")
        
        Returns:
            List of file paths
        """
        try:
            directory_path = Path(directory)
            return [str(f) for f in directory_path.glob(pattern) if f.is_file()]
        except Exception as e:
            return [f"Error listing files: {str(e)}"]
    
    @staticmethod
    async def read_code_files_async(
        directory: str,
        extensions: Optional[List[str]] = None,
        max_file_size: int = MAX_FILE_SIZE
    ) -> Dict[str, str]:
        """Async version of read_code_files for better performance.
        
        Args:
            directory: Path to the directory
            extensions: List of file extensions to include
            max_file_size: Maximum size per file in bytes
        
        Returns:
            Dictionary mapping file paths to their contents
        """
        if extensions is None:
            extensions = ['.py', '.java', '.js', '.ts', '.cs', '.cpp', '.c', '.h', '.go', '.rs']
        
        directory_path = Path(directory)
        if not directory_path.exists():
            return {"error": f"Directory {directory} does not exist"}
        
        async def read_file(file_path: Path) -> tuple:
            """Helper to read a single file async"""
            try:
                file_size = file_path.stat().st_size
                if file_size > max_file_size:
                    return (str(file_path), f"[Skipped - too large: {file_size / 1024 / 1024:.2f} MB]")
                
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    content = await f.read()
                    relative_path = file_path.relative_to(directory_path)
                    return (str(relative_path), content)
            except UnicodeDecodeError:
                return (str(file_path), "[Binary file - skipped]")
            except Exception as e:
                return (str(file_path), f"Error: {str(e)}")
        
        # Collect all files
        files_to_read = []
        for ext in extensions:
            files_to_read.extend(directory_path.rglob(f"*{ext}"))
        
        # Read all files concurrently
        tasks = [read_file(f) for f in files_to_read]
        results = await asyncio.gather(*tasks)
        
        return dict(results)


class CodeAnalysisTools:
    """Tools for analyzing code structure"""
    
    @staticmethod
    def extract_code_summary(
        code_files: Dict[str, str],
        max_lines_per_file: int = 100,
        use_ast_parsing: bool = True
    ) -> str:
        """Create an intelligent summary of code files using AST parsing.
        
        Args:
            code_files: Dictionary of file paths to contents
            max_lines_per_file: Maximum lines to include per file (fallback)
            use_ast_parsing: Use AST parsing for Python files (preserves key patterns)
        
        Returns:
            Formatted summary string
        """
        summary_parts = []
        summary_parts.append(f"# Code Repository Summary")
        summary_parts.append(f"Total files: {len(code_files)}\n")
        
        for file_path, content in code_files.items():
            # Skip metadata and error entries
            if file_path.startswith("_metadata") or "Error" in content[:50]:
                continue
            
            summary_parts.append(f"\n## File: {file_path}")
            lines = content.split('\n')
            summary_parts.append(f"Lines: {len(lines)}")
            
            # Try AST parsing for Python files
            if use_ast_parsing and file_path.endswith('.py'):
                try:
                    extracted = CodeAnalysisTools._extract_python_structure(content)
                    if extracted:
                        summary_parts.append("\n### Structure:")
                        summary_parts.append(extracted)
                        
                        # Add key code patterns
                        if len(lines) > max_lines_per_file:
                            summary_parts.append(f"\n### Key Code Patterns (intelligent extraction):")
                            key_lines = CodeAnalysisTools._extract_key_patterns(lines, max_lines_per_file)
                            summary_parts.append("```python")
                            summary_parts.append('\n'.join(key_lines))
                            summary_parts.append("```")
                        else:
                            summary_parts.append("\n### Full Content:")
                            summary_parts.append("```python")
                            summary_parts.append(content)
                            summary_parts.append("```")
                        continue
                except SyntaxError:
                    pass  # Fall through to simple truncation
            
            # Fallback: simple truncation for non-Python or unparseable files
            if len(lines) > max_lines_per_file:
                summary_parts.append(f"(Showing first {max_lines_per_file} lines)")
                truncated = lines[:max_lines_per_file]
            else:
                truncated = lines
            
            # Detect language by extension
            ext = Path(file_path).suffix
            lang_map = {'.py': 'python', '.js': 'javascript', '.ts': 'typescript',
                       '.java': 'java', '.cpp': 'cpp', '.c': 'c', '.go': 'go', '.rs': 'rust'}
            lang = lang_map.get(ext, '')
            
            summary_parts.append(f"```{lang}")
            summary_parts.append('\n'.join(truncated))
            summary_parts.append("```\n")
        
        return '\n'.join(summary_parts)
    
    @staticmethod
    def _extract_python_structure(code: str) -> Optional[str]:
        """Extract high-level structure from Python code using AST."""
        try:
            tree = ast.parse(code)
            structure = []
            
            for node in ast.walk(tree):
                if isinstance(node, ast.ClassDef):
                    methods = [m.name for m in node.body if isinstance(m, ast.FunctionDef)]
                    structure.append(f"- Class: {node.name}")
                    if methods:
                        structure.append(f"  Methods: {', '.join(methods)}")
                elif isinstance(node, ast.FunctionDef) and node.col_offset == 0:
                    # Top-level functions only
                    args = [arg.arg for arg in node.args.args]
                    structure.append(f"- Function: {node.name}({', '.join(args)})")
                elif isinstance(node, ast.Import):
                    names = [alias.name for alias in node.names]
                    structure.append(f"- Import: {', '.join(names)}")
                elif isinstance(node, ast.ImportFrom):
                    names = [alias.name for alias in node.names]
                    structure.append(f"- From {node.module} import: {', '.join(names)}")
            
            return '\n'.join(structure) if structure else None
        except SyntaxError:
            return None
    
    @staticmethod
    def _extract_key_patterns(lines: List[str], max_lines: int) -> List[str]:
        """Extract key patterns from code (imports, classes, functions, decorators)."""
        key_lines = []
        keywords = ['import ', 'from ', 'class ', 'def ', 'async def ', '@', 'if __name__']
        
        for line in lines:
            stripped = line.strip()
            # Include lines with key patterns
            if any(stripped.startswith(kw) for kw in keywords):
                key_lines.append(line)
            # Include significant comments
            elif stripped.startswith('#') and len(stripped) > 10:
                key_lines.append(line)
            # Include docstrings
            elif '"""' in stripped or "'''" in stripped:
                key_lines.append(line)
            
            if len(key_lines) >= max_lines:
                break
        
        # If not enough key patterns found, add first N lines
        if len(key_lines) < max_lines // 2:
            return lines[:max_lines]
        
        return key_lines
    
    @staticmethod
    def identify_file_types(code_files: Dict[str, str]) -> Dict[str, List[str]]:
        """
        Categorize files by their extension/type
        
        Args:
            code_files: Dictionary of file paths to contents
        
        Returns:
            Dictionary mapping extensions to list of files
        """
        file_types = {}
        
        for file_path in code_files.keys():
            ext = Path(file_path).suffix or "no_extension"
            if ext not in file_types:
                file_types[ext] = []
            file_types[ext].append(file_path)
        
        return file_types


class DocumentFormatter:
    """Tools for formatting and converting documents"""
    
    @staticmethod
    def markdown_to_docx(markdown_content: str, output_path: str) -> str:
        """Convert Markdown to DOCX format.
        
        Args:
            markdown_content: Markdown content
            output_path: Output DOCX file path
        
        Returns:
            Success message or error
        """
        # Check if required libraries are available
        if not DOCX_AVAILABLE:
            return "Error: python-docx not installed. Install with: pip install python-docx"
        if not MARKDOWN_AVAILABLE:
            return "Error: markdown not installed. Install with: pip install markdown"
        if not BS4_AVAILABLE:
            return "Error: beautifulsoup4 not installed. Install with: pip install beautifulsoup4"
        
        try:
            # Convert markdown to HTML
            if markdown_to_html is None:
                return "Error: markdown not installed. Install with: pip install markdown"
            if BeautifulSoup is None:
                return "Error: beautifulsoup4 not installed. Install with: pip install beautifulsoup4"
            if DocxDocument is None:
                return "Error: python-docx not installed. Install with: pip install python-docx"
            html = markdown_to_html(markdown_content)
            soup = BeautifulSoup(html, 'html.parser')
            
            # Create DOCX document
            doc = DocxDocument()
            
            # Simple conversion (can be enhanced)
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li', 'code', 'pre']):
                if element.name == 'h1':
                    doc.add_heading(element.get_text(), level=1)
                elif element.name == 'h2':
                    doc.add_heading(element.get_text(), level=2)
                elif element.name == 'h3':
                    doc.add_heading(element.get_text(), level=3)
                elif element.name in ['code', 'pre']:
                    # Add code blocks with monospace font
                    p = doc.add_paragraph(element.get_text())
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                else:
                    doc.add_paragraph(element.get_text())
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return f"DOCX document saved to {output_path}"
        
        except Exception as e:
            return f"Error converting to DOCX: {str(e)}"
    
    @staticmethod
    async def save_as_docx_async(markdown_content: str, output_path: str) -> str:
        """Async convert Markdown to DOCX format.
        
        This wraps the synchronous docx operations in asyncio.to_thread
        for non-blocking I/O in async contexts.
        
        Args:
            markdown_content: Markdown content
            output_path: Output DOCX file path
        
        Returns:
            Success message or error
        """
        # Run blocking docx operations in thread pool
        import asyncio
        return await asyncio.to_thread(
            FileTools.save_as_docx,
            markdown_content,
            output_path
        )
    
    @staticmethod
    def save_as_pdf(markdown_content: str, output_path: str) -> str:
        """Convert Markdown to PDF using WeasyPrint.
        
        Args:
            markdown_content: Markdown content
            output_path: Output PDF file path
        
        Returns:
            Success message or error
        """
        if not WEASYPRINT_AVAILABLE:
            return (
                "Error: weasyprint not installed.\n"
                "Install with: pip install weasyprint\n\n"
                "Alternative: Use pandoc command:\n"
                "  pandoc input.md -o output.pdf"
            )
        if not MARKDOWN_AVAILABLE:
            return "Error: markdown not installed. Install with: pip install markdown"
        if markdown_to_html is None:
            return "Error: markdown not installed. Install with: pip install markdown"
        if HTML is None:
            return (
                "Error: weasyprint not installed.\n"
                "Install with: pip install weasyprint\n\n"
                "Alternative: Use pandoc command:\n"
                "  pandoc input.md -o output.pdf"
            )
        
        try:
            # Convert markdown to HTML
            html_content = markdown_to_html(markdown_content, extensions=['extra', 'codehilite'])
            
            # Add basic CSS for better formatting
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        max-width: 800px;
                        margin: 40px auto;
                        padding: 20px;
                    }}
                    h1, h2, h3 {{ color: #333; }}
                    code {{
                        background: #f4f4f4;
                        padding: 2px 5px;
                        font-family: 'Courier New', monospace;
                    }}
                    pre {{
                        background: #f4f4f4;
                        padding: 10px;
                        overflow-x: auto;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                    }}
                    th, td {{
                        border: 1px solid #ddd;
                        padding: 8px;
                        text-align: left;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            # Convert HTML to PDF
            HTML(string=styled_html).write_pdf(output_path)
            return f"PDF document saved to {output_path}"
        
        except Exception as e:
            return f"Error converting to PDF: {str(e)}"
    
    @staticmethod
    async def save_as_pdf_async(markdown_content: str, output_path: str) -> str:
        """Async convert Markdown to PDF using WeasyPrint.
        
        This wraps the synchronous PDF operations in asyncio.to_thread
        for non-blocking I/O in async contexts.
        
        Args:
            markdown_content: Markdown content
            output_path: Output PDF file path
        
        Returns:
            Success message or error
        """
        # Run blocking PDF operations in thread pool
        import asyncio
        return await asyncio.to_thread(
            FileTools.save_as_pdf,
            markdown_content,
            output_path
        )
